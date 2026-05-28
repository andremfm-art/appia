"""
Moura IA — Assistente Multimodal
v3.0 — RAG corrigido, design refinado, branding Moura IA
"""

import streamlit as st

st.set_page_config(
    page_title="Moura IA Studio",
    page_icon="🟢",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={"About": "Moura IA v3.0 — Powered by NVIDIA NIM"},
)

# ─── CSS ─────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700;800&family=JetBrains+Mono:wght@400;500;600&display=swap');

*, *::before, *::after { box-sizing: border-box; }

:root {
    --green:   #22c55e;
    --green2:  #16a34a;
    --teal:    #2dd4bf;
    --dark:    #080b10;
    --surface: #0f1319;
    --card:    #141922;
    --border:  #1e2736;
    --border2: #273040;
    --text:    #e2e8f0;
    --muted:   #64748b;
    --muted2:  #94a3b8;
    --amber:   #f59e0b;
    --red:     #ef4444;
}

html, body { margin: 0; padding: 0; }

[data-testid="stAppViewContainer"] {
    background: var(--dark) !important;
    font-family: 'Outfit', sans-serif !important;
}

/* Sidebar */
[data-testid="stSidebar"] {
    background: var(--surface) !important;
    border-right: 1px solid var(--border) !important;
    padding-top: 0 !important;
}
[data-testid="stSidebar"] > div:first-child { padding-top: 0 !important; }
[data-testid="stSidebar"] * { font-family: 'Outfit', sans-serif !important; }
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span:not(.ctx-badge) { color: var(--muted2) !important; }
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3 { color: var(--text) !important; }

/* Main area */
[data-testid="stMain"] { background: var(--dark) !important; }

/* Chat messages */
.stChatMessage { background: transparent !important; border: none !important; gap: 12px !important; }

[data-testid="stChatMessageContent"] {
    background: var(--card) !important;
    border: 1px solid var(--border) !important;
    border-radius: 14px !important;
    padding: 18px 22px !important;
    font-family: 'Outfit', sans-serif !important;
    font-size: 0.95rem !important;
    font-weight: 400 !important;
    line-height: 1.75 !important;
    color: var(--text) !important;
    box-shadow: 0 2px 12px rgba(0,0,0,0.3) !important;
}

/* User bubble different style */
[data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatarUser"]) [data-testid="stChatMessageContent"] {
    background: #0f1f16 !important;
    border-color: rgba(34,197,94,0.25) !important;
}

[data-testid="stChatMessageContent"] p { color: var(--text) !important; }
[data-testid="stChatMessageContent"] strong { color: #fff !important; }
[data-testid="stChatMessageContent"] code {
    background: #0d1117 !important;
    color: var(--green) !important;
    padding: 2px 7px !important;
    border-radius: 5px !important;
    font-family: 'JetBrains Mono', monospace !important;
    font-size: 0.85em !important;
    border: 1px solid var(--border) !important;
}
[data-testid="stChatMessageContent"] pre {
    background: #0d1117 !important;
    border: 1px solid var(--border) !important;
    border-left: 3px solid var(--green) !important;
    border-radius: 10px !important;
    padding: 18px !important;
    overflow-x: auto !important;
}
[data-testid="stChatMessageContent"] pre code {
    background: transparent !important;
    border: none !important;
    padding: 0 !important;
    color: #e2e8f0 !important;
}
[data-testid="stChatMessageContent"] blockquote {
    border-left: 3px solid var(--green) !important;
    padding-left: 14px !important;
    color: var(--muted2) !important;
    margin: 8px 0 !important;
}
[data-testid="stChatMessageContent"] table {
    width: 100% !important;
    border-collapse: collapse !important;
}
[data-testid="stChatMessageContent"] th {
    background: var(--border) !important;
    color: var(--green) !important;
    padding: 8px 12px !important;
    font-weight: 600 !important;
    text-align: left !important;
}
[data-testid="stChatMessageContent"] td {
    padding: 7px 12px !important;
    border-bottom: 1px solid var(--border) !important;
    color: var(--text) !important;
}

/* Chat input */
[data-testid="stChatInput"] {
    background: var(--card) !important;
    border: 1.5px solid var(--border2) !important;
    border-radius: 16px !important;
    box-shadow: 0 0 0 0 transparent !important;
    transition: border-color 0.2s !important;
}
[data-testid="stChatInput"]:focus-within {
    border-color: var(--green) !important;
    box-shadow: 0 0 0 3px rgba(34,197,94,0.1) !important;
}
[data-testid="stChatInput"] textarea {
    background: transparent !important;
    color: var(--text) !important;
    font-family: 'Outfit', sans-serif !important;
    font-size: 0.95rem !important;
}
[data-testid="stChatInput"] button {
    background: var(--green) !important;
    border-radius: 10px !important;
}

/* Buttons */
.stButton > button {
    background: var(--card) !important;
    color: var(--muted2) !important;
    border: 1px solid var(--border) !important;
    border-radius: 9px !important;
    font-family: 'Outfit', sans-serif !important;
    font-weight: 500 !important;
    font-size: 0.88rem !important;
    transition: all 0.18s ease !important;
    padding: 6px 14px !important;
}
.stButton > button:hover {
    border-color: var(--green) !important;
    color: var(--green) !important;
    background: rgba(34,197,94,0.06) !important;
    transform: translateY(-1px) !important;
}
.stButton > button:active { transform: translateY(0) !important; }

/* Primary action button */
.btn-primary > button {
    background: linear-gradient(135deg, var(--green2), var(--green)) !important;
    color: #fff !important;
    border: none !important;
    font-weight: 600 !important;
}
.btn-primary > button:hover {
    background: linear-gradient(135deg, #15803d, var(--green2)) !important;
    color: #fff !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 4px 14px rgba(34,197,94,0.3) !important;
}

/* Download buttons */
.stDownloadButton > button {
    background: var(--card) !important;
    color: var(--muted2) !important;
    border: 1px solid var(--border) !important;
    border-radius: 8px !important;
    font-size: 0.8rem !important;
    padding: 5px 12px !important;
    font-family: 'JetBrains Mono', monospace !important;
}
.stDownloadButton > button:hover {
    border-color: var(--teal) !important;
    color: var(--teal) !important;
    background: rgba(45,212,191,0.06) !important;
}

/* Selectbox */
.stSelectbox > div > div {
    background: var(--card) !important;
    border-color: var(--border) !important;
    color: var(--text) !important;
    border-radius: 9px !important;
    font-family: 'Outfit', sans-serif !important;
}
.stSelectbox [data-baseweb="select"] { background: var(--card) !important; }
.stSelectbox svg { color: var(--muted) !important; }

/* Sliders */
.stSlider [data-baseweb="slider"] { padding: 0 !important; }
.stSlider [data-testid="stTickBar"] { display: none !important; }

/* Checkboxes */
.stCheckbox label { color: var(--muted2) !important; font-size: 0.88rem !important; }
.stCheckbox [data-baseweb="checkbox"] div { border-color: var(--border2) !important; }

/* File uploader */
[data-testid="stFileUploader"] {
    background: var(--card) !important;
    border: 1.5px dashed var(--border2) !important;
    border-radius: 12px !important;
}
[data-testid="stFileUploader"] label { color: var(--muted2) !important; }
[data-testid="stFileUploader"] button {
    background: var(--card) !important;
    border-color: var(--border2) !important;
    color: var(--muted2) !important;
}

/* Alerts */
.stSuccess, .stError, .stWarning, .stInfo {
    border-radius: 10px !important;
    font-family: 'Outfit', sans-serif !important;
    font-size: 0.88rem !important;
}

/* Progress */
.stProgress > div > div { background: var(--green) !important; border-radius: 99px !important; }

/* Divider */
hr { border-color: var(--border) !important; margin: 16px 0 !important; }

/* Spinner */
.stSpinner > div { color: var(--green) !important; }

/* ── Custom components ─────────────────────────────────────── */
.moura-sidebar-header {
    background: linear-gradient(180deg, #0a1a0f 0%, var(--surface) 100%);
    padding: 24px 20px 20px;
    border-bottom: 1px solid var(--border);
    margin-bottom: 8px;
}
.moura-logo {
    font-family: 'Outfit', sans-serif;
    font-weight: 800;
    font-size: 1.5rem;
    letter-spacing: -0.5px;
    background: linear-gradient(135deg, #22c55e, #2dd4bf);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    line-height: 1;
}
.moura-tagline {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.62rem;
    color: var(--muted);
    letter-spacing: 2px;
    text-transform: uppercase;
    margin-top: 4px;
}

.section-label {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.65rem;
    font-weight: 600;
    letter-spacing: 2px;
    text-transform: uppercase;
    color: var(--muted) !important;
    margin: 16px 0 8px;
    padding: 0 2px;
}

.rag-stat {
    display: flex;
    align-items: center;
    justify-content: space-between;
    background: rgba(34,197,94,0.06);
    border: 1px solid rgba(34,197,94,0.2);
    border-radius: 10px;
    padding: 10px 14px;
    margin: 6px 0;
}
.rag-stat-label { font-size: 0.75rem; color: var(--muted); font-family: 'JetBrains Mono', monospace; }
.rag-stat-val { font-size: 0.9rem; color: var(--green); font-weight: 700; font-family: 'Outfit', sans-serif; }

.source-chip {
    display: inline-flex;
    align-items: center;
    gap: 5px;
    background: rgba(34,197,94,0.08);
    border: 1px solid rgba(34,197,94,0.2);
    border-radius: 99px;
    padding: 3px 10px;
    font-size: 0.72rem;
    color: var(--green);
    font-family: 'JetBrains Mono', monospace;
    margin: 2px 2px;
    max-width: 180px;
    overflow: hidden;
    text-overflow: ellipsis;
    white-space: nowrap;
}

.conv-btn {
    display: flex;
    align-items: center;
    gap: 8px;
    padding: 8px 12px;
    border-radius: 9px;
    font-size: 0.84rem;
    color: var(--muted2);
    cursor: pointer;
    border: 1px solid transparent;
    transition: all 0.15s;
    width: 100%;
    overflow: hidden;
    white-space: nowrap;
    text-overflow: ellipsis;
    background: transparent;
    margin: 1px 0;
    font-family: 'Outfit', sans-serif;
}
.conv-btn:hover { background: var(--card); border-color: var(--border); }
.conv-btn.active {
    background: rgba(34,197,94,0.08);
    border-color: rgba(34,197,94,0.3);
    color: var(--green);
}

/* Main header */
.main-header {
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 20px 0 16px;
    border-bottom: 1px solid var(--border);
    margin-bottom: 28px;
}
.main-title {
    font-family: 'Outfit', sans-serif;
    font-weight: 800;
    font-size: 1.5rem;
    letter-spacing: -0.5px;
    background: linear-gradient(135deg, #22c55e, #2dd4bf);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
}
.badge {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.62rem;
    font-weight: 600;
    letter-spacing: 1.5px;
    padding: 4px 10px;
    border-radius: 99px;
    display: inline-flex;
    align-items: center;
    gap: 5px;
}
.badge-green  { background: rgba(34,197,94,0.12); color: var(--green); border: 1px solid rgba(34,197,94,0.25); }
.badge-teal   { background: rgba(45,212,191,0.12); color: var(--teal); border: 1px solid rgba(45,212,191,0.25); }
.badge-amber  { background: rgba(245,158,11,0.12); color: var(--amber); border: 1px solid rgba(245,158,11,0.25); }
.badge-muted  { background: rgba(100,116,139,0.12); color: var(--muted2); border: 1px solid rgba(100,116,139,0.2); }

/* Context indicator above assistant reply */
.ctx-row { display: flex; flex-wrap: wrap; gap: 6px; margin-bottom: 10px; }

/* Model tag */
.model-footer {
    display: flex;
    align-items: center;
    gap: 10px;
    margin-top: 12px;
    padding-top: 10px;
    border-top: 1px solid var(--border);
}
.model-footer-tag {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.67rem;
    color: var(--muted);
    display: inline-flex;
    align-items: center;
    gap: 5px;
}

/* Welcome screen */
.welcome-wrap {
    text-align: center;
    padding: 64px 0 32px;
    max-width: 560px;
    margin: 0 auto;
}
.welcome-icon { font-size: 3.5rem; margin-bottom: 20px; line-height: 1; }
.welcome-title {
    font-family: 'Outfit', sans-serif;
    font-weight: 800;
    font-size: 2rem;
    letter-spacing: -1px;
    background: linear-gradient(135deg, #22c55e, #2dd4bf);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    margin-bottom: 10px;
}
.welcome-sub { font-size: 0.9rem; color: var(--muted2); line-height: 1.6; }
.caps-grid {
    display: grid;
    grid-template-columns: repeat(2, 1fr);
    gap: 12px;
    margin: 32px auto 0;
    max-width: 520px;
}
.cap-card {
    background: var(--card);
    border: 1px solid var(--border);
    border-radius: 14px;
    padding: 18px;
    text-align: left;
    transition: border-color 0.2s;
}
.cap-card:hover { border-color: var(--border2); }
.cap-icon { font-size: 1.3rem; margin-bottom: 8px; }
.cap-title { font-weight: 600; font-size: 0.9rem; color: var(--text); margin-bottom: 4px; }
.cap-desc { font-size: 0.8rem; color: var(--muted2); line-height: 1.5; }

/* RAG debug expander */
.rag-debug {
    background: #0a1a0f;
    border: 1px solid rgba(34,197,94,0.2);
    border-radius: 10px;
    padding: 12px 16px;
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.72rem;
    color: var(--muted2);
    margin-top: 8px;
    max-height: 200px;
    overflow-y: auto;
    white-space: pre-wrap;
    word-break: break-word;
}

/* Scrollbar */
::-webkit-scrollbar { width: 4px; height: 4px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: var(--border2); border-radius: 99px; }
::-webkit-scrollbar-thumb:hover { background: var(--muted); }
</style>
""", unsafe_allow_html=True)

# ─── IMPORTS ─────────────────────────────────────────────────────────────────
import os, base64, hashlib, logging, pickle, sqlite3, atexit, re, time
from io import BytesIO
from datetime import datetime
from typing import Optional, List, Dict, Tuple

import numpy as np
from openai import OpenAI
from pypdf import PdfReader
from gtts import gTTS
from docx import Document
from PIL import Image

try:
    from duckduckgo_search import DDGS
    DDGS_OK = True
except ImportError:
    DDGS_OK = False

try:
    from sentence_transformers import SentenceTransformer
    ST_OK = True
except ImportError:
    ST_OK = False

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger("moura_ia")

for _d in ["data", "data/images", "data/tts", "data/exports"]:
    os.makedirs(_d, exist_ok=True)

# ─── MODELOS ─────────────────────────────────────────────────────────────────
TEXT_MODELS = {
    "Llama 3.1 70B":  "meta/llama-3.1-70b-instruct",
    "Llama 3.1 8B":   "meta/llama-3.1-8b-instruct",
    "Llama 3.3 70B":  "meta/llama-3.3-70b-instruct",
    "Nemotron 70B":   "nvidia/nemotron-70b-instruct",
    "Mistral NeMo":   "nv-mistralai/mistral-nemo-12b-instruct",
}
VISION_MODELS = [
    "meta/llama-3.2-11b-vision-instruct",
    "meta/llama-3.2-90b-vision-instruct",
]
FALLBACK = "meta/llama-3.1-8b-instruct"
EMBED_MODEL = "all-MiniLM-L6-v2"
VECTORS_FILE = "data/vectors.pkl"
CHUNK_SIZE = 600
CHUNK_OVERLAP = 100
TOP_K = 5
SIM_THRESHOLD = 0.20

# ─── API KEY ─────────────────────────────────────────────────────────────────
def get_api_key() -> Optional[str]:
    try:
        if hasattr(st, "secrets") and "NVIDIA_API_KEY" in st.secrets:
            k = st.secrets["NVIDIA_API_KEY"]
            if k and k.startswith("nvapi-"): return k
    except: pass
    k = os.environ.get("NVIDIA_API_KEY", "")
    if k.startswith("nvapi-"): return k
    return st.session_state.get("api_key")

def show_api_screen():
    st.markdown("""
    <div style='max-width:440px;margin:100px auto 0;text-align:center'>
        <div style='font-family:Outfit,sans-serif;font-weight:800;font-size:2.2rem;
             background:linear-gradient(135deg,#22c55e,#2dd4bf);
             -webkit-background-clip:text;-webkit-text-fill-color:transparent;
             margin-bottom:6px'>Moura IA</div>
        <div style='color:#64748b;font-size:0.88rem;margin-bottom:36px'>
             Configure sua NVIDIA API Key para começar
        </div>
    </div>""", unsafe_allow_html=True)
    col = st.columns([1,2,1])[1]
    with col:
        k = st.text_input("NVIDIA API Key", type="password", placeholder="nvapi-...")
        up = st.file_uploader("Ou arquivo .env / .txt", type=["env","txt","key"])
        if up:
            c = up.read().decode(errors="ignore")
            for ln in c.splitlines():
                ln = ln.strip()
                if ln.startswith("NVIDIA_API_KEY="):
                    k = ln.split("=",1)[1].strip().strip("\"'")
                elif ln.startswith("nvapi-"):
                    k = ln
                if k and k.startswith("nvapi-"): break
        if st.button("Entrar →", use_container_width=True):
            if k and k.startswith("nvapi-"):
                st.session_state.api_key = k; st.rerun()
            else:
                st.error("Chave inválida — deve começar com `nvapi-`")
        st.markdown("<div style='text-align:center;margin-top:12px'><a href='https://build.nvidia.com' target='_blank' style='color:#22c55e;font-size:0.8rem'>🔗 Obter chave gratuita</a></div>", unsafe_allow_html=True)

# ─── DATABASE ─────────────────────────────────────────────────────────────────
class DB:
    def __init__(self, path="data/chat.db"):
        self.path = path
        self._c: Optional[sqlite3.Connection] = None
        self._init()

    def _connect(self):
        c = sqlite3.connect(self.path, check_same_thread=False)
        c.row_factory = sqlite3.Row
        c.execute("PRAGMA journal_mode=WAL")
        c.execute("PRAGMA foreign_keys=ON")
        return c

    def _init(self):
        c = self._connect()
        c.executescript("""
            CREATE TABLE IF NOT EXISTS conversations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT DEFAULT 'Nova Conversa',
                created_at TEXT DEFAULT (datetime('now','localtime')),
                updated_at TEXT DEFAULT (datetime('now','localtime'))
            );
            CREATE TABLE IF NOT EXISTS messages (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                conversation_id INTEGER NOT NULL,
                role TEXT NOT NULL,
                content TEXT NOT NULL,
                image_path TEXT,
                model TEXT,
                ts TEXT DEFAULT (datetime('now','localtime')),
                FOREIGN KEY(conversation_id) REFERENCES conversations(id) ON DELETE CASCADE
            );
            CREATE INDEX IF NOT EXISTS idx_msg_conv ON messages(conversation_id);
        """)
        c.commit(); c.close()

    @property
    def conn(self):
        if self._c is None: self._c = self._connect()
        return self._c

    def ex(self, sql, p=()):
        try:
            cur = self.conn.execute(sql, p); self.conn.commit(); return cur
        except:
            self._c = self._connect()
            cur = self.conn.execute(sql, p); self.conn.commit(); return cur

    def all(self, sql, p=()): return self.ex(sql,p).fetchall()
    def one(self, sql, p=()): return self.ex(sql,p).fetchone()
    def close(self):
        if self._c: self._c.close(); self._c = None

# ─── EMBEDDING ────────────────────────────────────────────────────────────────
@st.cache_resource(show_spinner=False)
def load_embedder():
    if not ST_OK: return None, 384
    log.info("Carregando embedder...")
    m = SentenceTransformer(EMBED_MODEL)
    return m, m.get_sentence_embedding_dimension()

# ─── VECTOR STORE ─────────────────────────────────────────────────────────────
def _empty_store():
    return {"docs": [], "embs": [], "ids": [], "srcs": []}

def load_vs() -> dict:
    if not os.path.exists(VECTORS_FILE): return _empty_store()
    try:
        with open(VECTORS_FILE,"rb") as f: data = pickle.load(f)
        if not all(k in data for k in ("docs","embs","ids","srcs")):
            return _empty_store()
        return data
    except:
        return _empty_store()

def save_vs(data):
    with open(VECTORS_FILE,"wb") as f: pickle.dump(data,f)

def cosine(a, b):
    a,b = np.asarray(a,dtype=np.float32), np.asarray(b,dtype=np.float32)
    na,nb = np.linalg.norm(a), np.linalg.norm(b)
    if na==0 or nb==0: return 0.0
    return float(np.dot(a,b)/(na*nb))

def clean(t): return re.sub(r"\s+"," ",t).strip()

def chunks(text, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    """Chunk text preserving sentence boundaries when possible."""
    text = clean(text)
    if len(text) <= size: return [text] if text else []
    result = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        # Try to break at sentence boundary
        if end < len(text):
            for sep in [". ", ".\n", "! ", "? ", "\n\n", "\n"]:
                idx = text.rfind(sep, start + size//2, end)
                if idx != -1:
                    end = idx + len(sep)
                    break
        chunk = text[start:end].strip()
        if len(chunk) > 40:
            result.append(chunk)
        start = end - overlap
        if start >= len(text): break
    return result

# ─── RAG SERVICE ──────────────────────────────────────────────────────────────
class RAG:
    def __init__(self):
        self.embedder, self.dim = load_embedder()
        self.vs = load_vs()
        self._check_dim()

    def _check_dim(self):
        if self.vs["embs"] and self.embedder:
            if len(self.vs["embs"][0]) != self.dim:
                log.warning("Dimensão incompatível — recriando store")
                self.vs = _empty_store(); save_vs(self.vs)

    def embed(self, text: str) -> Optional[np.ndarray]:
        if not self.embedder: return None
        try: return self.embedder.encode(text, normalize_embeddings=True)
        except Exception as e: log.error(f"Embed erro: {e}"); return None

    def add(self, filename: str, text: str) -> int:
        text = clean(text)
        if not text: return 0
        cks = chunks(text)
        added = 0
        for i, ck in enumerate(cks):
            emb = self.embed(ck)
            if emb is None: continue
            if self.vs["embs"] and len(emb) != len(self.vs["embs"][0]):
                continue
            uid = f"{filename}|{i}|{hashlib.md5(ck.encode()).hexdigest()[:8]}"
            if uid in self.vs["ids"]: continue
            self.vs["docs"].append(ck)
            self.vs["embs"].append(emb.tolist())
            self.vs["ids"].append(uid)
            self.vs["srcs"].append(filename)
            added += 1
        save_vs(self.vs)
        return added

    def search(self, query: str, k=TOP_K, threshold=SIM_THRESHOLD) -> Tuple[str, List[dict]]:
        """
        Retorna (contexto_formatado, lista_de_resultados_para_debug)
        CORREÇÃO CRÍTICA: usa todos os embeddings disponíveis, 
        threshold mais baixo, e retorna contexto mais rico.
        """
        if not self.vs["embs"]:
            return "", []
        
        emb = self.embed(query)
        if emb is None:
            return "", []

        q = np.array(emb, dtype=np.float32)
        
        # Calcular similaridade com TODOS os chunks
        scores = []
        for i, e in enumerate(self.vs["embs"]):
            sim = cosine(q, np.array(e, dtype=np.float32))
            scores.append((sim, i))
        
        scores.sort(key=lambda x: x[0], reverse=True)
        
        # Pegar top-k sem threshold restritivo primeiro
        top_k = scores[:k]
        
        # Filtrar por threshold mínimo
        filtered = [(s,i) for s,i in top_k if s >= threshold]
        
        # Se nenhum passar o threshold mas temos documentos,
        # usar os top-3 mesmo assim (o modelo decide se é relevante)
        if not filtered and scores:
            filtered = scores[:min(3, len(scores))]
        
        if not filtered:
            return "", []
        
        results = []
        parts = []
        for rank, (sim, idx) in enumerate(filtered):
            src = self.vs["srcs"][idx] if idx < len(self.vs["srcs"]) else "doc"
            doc = self.vs["docs"][idx]
            results.append({"rank": rank+1, "sim": sim, "src": src, "text": doc})
            parts.append(f"[Fonte: {src} | Relevância: {sim:.2f}]\n{doc}")
        
        ctx = "\n\n---\n\n".join(parts)
        return ctx, results

    def clear(self):
        self.vs = _empty_store(); save_vs(self.vs)

    @property
    def count(self): return len(self.vs["docs"])
    
    @property
    def sources(self): return list(dict.fromkeys(self.vs.get("srcs",[])))

# ─── WEB SEARCH ───────────────────────────────────────────────────────────────
def web_search(query: str, n=4) -> str:
    if not DDGS_OK: return ""
    try:
        with DDGS() as d:
            res = list(d.text(query, max_results=n))
        if not res: return ""
        out = []
        for r in res:
            title = r.get("title","")
            body  = r.get("body","")[:400]
            href  = r.get("href","")
            out.append(f"• {title}\n  {body}\n  {href}")
        return "\n\n".join(out)
    except Exception as e:
        log.warning(f"Web search erro: {e}"); return ""

# ─── LLM SERVICE ──────────────────────────────────────────────────────────────
class LLM:
    def __init__(self, key):
        self.client = OpenAI(api_key=key, base_url="https://integrate.api.nvidia.com/v1")

    def _call(self, msgs, model, temp, toks, stream=False):
        return self.client.chat.completions.create(
            model=model, messages=msgs, temperature=temp,
            max_tokens=toks, stream=stream, timeout=60
        )

    def gen(self, msgs, model, temp=0.7, toks=1024) -> Tuple[str,str]:
        for m in ([model, FALLBACK] if model != FALLBACK else [model]):
            try:
                r = self._call(msgs, m, temp, toks)
                return r.choices[0].message.content or "", m
            except Exception as e:
                log.warning(f"Model {m}: {e}")
        return "❌ Erro ao gerar resposta.", "error"

    def stream(self, msgs, model, temp=0.7, toks=1024):
        try:
            for chunk in self._call(msgs, model, temp, toks, stream=True):
                d = chunk.choices[0].delta
                if d and d.content: yield d.content
        except Exception as e:
            log.warning(f"Stream erro: {e}")
            text, _ = self.gen(msgs, model, temp, toks)
            yield text

    def vision(self, b64, prompt, mime="image/jpeg", temp=0.7, toks=1024) -> Tuple[str,str]:
        msgs = [{"role":"user","content":[
            {"type":"image_url","image_url":{"url":f"data:{mime};base64,{b64}"}},
            {"type":"text","text":prompt}
        ]}]
        for vm in VISION_MODELS:
            try:
                r = self._call(msgs, vm, temp, toks)
                return r.choices[0].message.content or "", vm
            except Exception as e:
                log.warning(f"Vision {vm}: {e}")
        return "❌ Erro no modelo de visão.", "error"

# ─── TTS ──────────────────────────────────────────────────────────────────────
def tts(text: str, lang="pt") -> Optional[str]:
    try:
        c = re.sub(r"[*_`#\[\](){}<>~]","", text)
        c = re.sub(r"https?://\S+","", c).strip()
        if not c: return None
        h = hashlib.md5(c[:800].encode()).hexdigest()
        p = f"data/tts/{h}.mp3"
        if not os.path.exists(p):
            gTTS(text=c[:5000], lang=lang, slow=False).save(p)
        return p
    except Exception as e:
        log.error(f"TTS: {e}"); return None

# ─── EXPORT ───────────────────────────────────────────────────────────────────
def export(msgs: List[Dict], fmt="txt") -> bytes:
    if not msgs: return b""
    try:
        if fmt=="txt":
            return "\n\n".join(
                f"[{'Usuário' if m['role']=='user' else 'Moura IA'}]\n{m['content']}"
                for m in msgs
            ).encode("utf-8")
        if fmt=="md":
            lines = ["# Conversa — Moura IA\n"]
            for m in msgs:
                r = "👤 Usuário" if m["role"]=="user" else "🤖 Moura IA"
                lines.append(f"## {r}\n{m['content']}\n")
            return "\n".join(lines).encode("utf-8")
        if fmt=="docx":
            doc = Document()
            doc.add_heading("Conversa — Moura IA", 1)
            for m in msgs:
                doc.add_heading("Usuário" if m["role"]=="user" else "Moura IA", 2)
                doc.add_paragraph(m["content"]); doc.add_paragraph("")
            buf = BytesIO(); doc.save(buf); return buf.getvalue()
    except Exception as e:
        log.error(f"Export {fmt}: {e}")
    return b""

# ─── SYSTEM PROMPT ────────────────────────────────────────────────────────────
def build_system(rag_ctx: str, web_ctx: str) -> str:
    base = (
        "Você é Moura IA, um assistente inteligente, preciso e multimodal. "
        "Responda sempre em português brasileiro de forma clara e bem formatada. "
        "Use Markdown quando útil (listas, negrito, código, tabelas). "
        "Seja objetivo mas completo."
    )
    if rag_ctx:
        base += (
            "\n\n━━━ CONTEXTO DOS DOCUMENTOS ━━━\n"
            "Use as informações abaixo como fonte primária para responder. "
            "Se a resposta estiver nos documentos, cite qual documento. "
            "Se não estiver, diga claramente que a informação não está nos documentos carregados.\n\n"
            + rag_ctx
        )
    if web_ctx:
        base += (
            "\n\n━━━ INFORMAÇÕES DA WEB ━━━\n"
            "Use como complemento se os documentos não contiverem a resposta:\n\n"
            + web_ctx
        )
    if not rag_ctx and not web_ctx:
        base += (
            "\n\nResponda com base no seu conhecimento geral. "
            "Seja honesto quando não souber algo."
        )
    return base

def fmt_model(mid): return mid.split("/")[-1] if "/" in mid else mid

# ═══════════════════════════════════════════════════════════════════════════════
#  BOOT
# ═══════════════════════════════════════════════════════════════════════════════
API_KEY = get_api_key()
if not API_KEY:
    show_api_screen(); st.stop()

# Singleton services
if "db"  not in st.session_state: st.session_state.db  = DB()
if "rag" not in st.session_state: st.session_state.rag = RAG()
if "llm" not in st.session_state: st.session_state.llm = LLM(API_KEY)

db:  DB  = st.session_state.db
rag: RAG = st.session_state.rag
llm: LLM = st.session_state.llm

# Session defaults
def _ss(k, v):
    if k not in st.session_state: st.session_state[k] = v

_ss("conv_id",    None)
_ss("messages",   None)
_ss("streaming",  True)
_ss("web_on",     True)
_ss("tts_on",     False)
_ss("temp",       0.7)
_ss("maxtok",     1024)
_ss("model_name", list(TEXT_MODELS.keys())[0])
_ss("img_key",    0)
_ss("show_debug", False)

# Load or create conversation
if st.session_state.conv_id is None:
    rows = db.all("SELECT id FROM conversations ORDER BY updated_at DESC LIMIT 1")
    if rows:
        st.session_state.conv_id = rows[0]["id"]
    else:
        cur = db.ex("INSERT INTO conversations (title) VALUES ('Nova Conversa')")
        st.session_state.conv_id = cur.lastrowid

if st.session_state.messages is None:
    rows = db.all(
        "SELECT role,content,image_path FROM messages WHERE conversation_id=? ORDER BY id",
        (st.session_state.conv_id,)
    )
    st.session_state.messages = []
    for r in rows:
        m = {"role": r["role"], "content": r["content"]}
        if r["image_path"]: m["image"] = r["image_path"]
        st.session_state.messages.append(m)

# ═══════════════════════════════════════════════════════════════════════════════
#  SIDEBAR
# ═══════════════════════════════════════════════════════════════════════════════
with st.sidebar:

    # Header
    st.markdown("""
    <div class='moura-sidebar-header'>
        <div class='moura-logo'>Moura IA</div>
        <div class='moura-tagline'>Assistente Multimodal · NVIDIA NIM</div>
    </div>""", unsafe_allow_html=True)

    # ── Conversas ────────────────────────────────────────────────────────────
    st.markdown("<div class='section-label'>Conversas</div>", unsafe_allow_html=True)

    if st.button("＋  Nova conversa", use_container_width=True):
        cur = db.ex("INSERT INTO conversations (title) VALUES ('Nova Conversa')")
        st.session_state.conv_id = cur.lastrowid
        st.session_state.messages = []
        st.session_state.img_key += 1
        st.rerun()

    convs = db.all("SELECT id,title FROM conversations ORDER BY updated_at DESC LIMIT 20")
    for c in convs:
        active = c["id"] == st.session_state.conv_id
        label  = (c["title"] or "Nova Conversa")[:38]
        icon   = "▶ " if active else "   "
        if st.button(f"{icon}{label}", key=f"c_{c['id']}", use_container_width=True):
            if c["id"] != st.session_state.conv_id:
                st.session_state.conv_id = c["id"]
                rows = db.all(
                    "SELECT role,content,image_path FROM messages WHERE conversation_id=? ORDER BY id",
                    (c["id"],)
                )
                st.session_state.messages = []
                for r in rows:
                    msg = {"role": r["role"], "content": r["content"]}
                    if r["image_path"]: msg["image"] = r["image_path"]
                    st.session_state.messages.append(msg)
                st.session_state.img_key += 1
                st.rerun()

    st.divider()

    # ── Modelo ───────────────────────────────────────────────────────────────
    st.markdown("<div class='section-label'>Modelo</div>", unsafe_allow_html=True)
    idx = list(TEXT_MODELS.keys()).index(st.session_state.model_name) \
          if st.session_state.model_name in TEXT_MODELS else 0
    mn = st.selectbox("", list(TEXT_MODELS.keys()), index=idx, label_visibility="collapsed")
    st.session_state.model_name = mn
    MODEL_ID = TEXT_MODELS[mn]

    st.markdown("<div class='section-label'>Parâmetros</div>", unsafe_allow_html=True)
    st.session_state.temp   = st.slider("Temperatura", 0.0, 1.0, st.session_state.temp, 0.05)
    st.session_state.maxtok = st.slider("Max tokens",  256, 4096, st.session_state.maxtok, 128)

    c1, c2 = st.columns(2)
    with c1: st.session_state.streaming = st.checkbox("Streaming", st.session_state.streaming)
    with c2: st.session_state.tts_on    = st.checkbox("Voz (TTS)", st.session_state.tts_on)
    st.session_state.web_on = st.checkbox("🌐 Busca Web", st.session_state.web_on)
    st.session_state.show_debug = st.checkbox("🔍 Debug RAG", st.session_state.show_debug)

    st.divider()

    # ── Documentos RAG ───────────────────────────────────────────────────────
    st.markdown("<div class='section-label'>Documentos (RAG)</div>", unsafe_allow_html=True)

    if rag.count > 0:
        st.markdown(f"""
        <div class='rag-stat'>
            <span class='rag-stat-label'>CHUNKS INDEXADOS</span>
            <span class='rag-stat-val'>{rag.count}</span>
        </div>
        <div style='margin:6px 0 10px;display:flex;flex-wrap:wrap;gap:4px'>
        {"".join(f"<span class='source-chip'>📄 {s[:24]}</span>" for s in rag.sources[:6])}
        </div>""", unsafe_allow_html=True)

    pdfs = st.file_uploader(
        "Carregar PDFs",
        type=["pdf"],
        accept_multiple_files=True,
        label_visibility="collapsed",
        key=f"pdf_up_{st.session_state.img_key}"
    )

    if pdfs:
        with st.container():
            st.markdown("<div class='btn-primary'>", unsafe_allow_html=True)
            do_index = st.button("⚡ Indexar PDFs", use_container_width=True)
            st.markdown("</div>", unsafe_allow_html=True)
        if do_index:
            total, errs = 0, []
            prog = st.progress(0)
            for i, f in enumerate(pdfs):
                try:
                    reader = PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        t = page.extract_text()
                        if t: text += t + "\n"
                    if not text.strip():
                        errs.append(f"'{f.name}': sem texto")
                        continue
                    n = rag.add(f.name, text)
                    total += n
                    log.info(f"Indexado {f.name}: {n} chunks")
                except Exception as e:
                    errs.append(f"'{f.name}': {e}")
                prog.progress((i+1)/len(pdfs))
            prog.empty()
            if total > 0:
                st.success(f"✅ {total} chunks indexados em {len(pdfs)} arquivo(s)")
                st.rerun()
            for e in errs: st.warning(e)

    if rag.count > 0:
        if st.button("🗑️ Limpar documentos", use_container_width=True):
            rag.clear(); st.success("Documentos removidos"); st.rerun()

    st.divider()

    # ── Export & ações ───────────────────────────────────────────────────────
    st.markdown("<div class='section-label'>Exportar conversa</div>", unsafe_allow_html=True)
    exp_msgs = [m for m in st.session_state.messages if m["role"] in ("user","assistant")]
    c1,c2,c3 = st.columns(3)
    c1.download_button("TXT",  export(exp_msgs,"txt"),  "conversa.txt",  "text/plain",            use_container_width=True)
    c2.download_button("MD",   export(exp_msgs,"md"),   "conversa.md",   "text/markdown",          use_container_width=True)
    c3.download_button("DOCX", export(exp_msgs,"docx"), "conversa.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

    if st.button("🗑️ Limpar conversa", use_container_width=True):
        db.ex("DELETE FROM messages WHERE conversation_id=?", (st.session_state.conv_id,))
        st.session_state.messages = []
        st.session_state.img_key += 1
        st.rerun()

    if st.button("🔑 Trocar API Key", use_container_width=True):
        for k in ["api_key","llm"]: st.session_state.pop(k, None)
        st.rerun()

# ═══════════════════════════════════════════════════════════════════════════════
#  ÁREA PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════════════
MODEL_ID = TEXT_MODELS[st.session_state.model_name]

# Header
rag_badge = f"<span class='badge badge-green'>📚 {rag.count} chunks</span>" if rag.count>0 else ""
web_badge = "<span class='badge badge-teal'>🌐 Web ON</span>" if st.session_state.web_on else ""
mod_badge = f"<span class='badge badge-muted'>{fmt_model(MODEL_ID)}</span>"

st.markdown(f"""
<div class='main-header'>
    <div class='main-title'>Moura IA</div>
    <div style='display:flex;gap:6px;flex-wrap:wrap;align-items:center'>
        {rag_badge}{web_badge}{mod_badge}
    </div>
</div>""", unsafe_allow_html=True)

# ── Histórico ────────────────────────────────────────────────────────────────
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        if msg.get("image") and os.path.exists(msg["image"]):
            st.image(msg["image"], width=380)
        st.markdown(msg["content"])

# ── Welcome screen ───────────────────────────────────────────────────────────
if not st.session_state.messages:
    st.markdown("""
    <div class='welcome-wrap'>
        <div class='welcome-icon'>🟢</div>
        <div class='welcome-title'>Moura IA</div>
        <div class='welcome-sub'>Assistente multimodal com RAG, busca web e visão computacional.<br>
        Carregue documentos, envie imagens ou pergunte qualquer coisa.</div>
    </div>
    <div class='caps-grid'>
        <div class='cap-card'><div class='cap-icon'>📄</div>
            <div class='cap-title'>Análise de PDFs</div>
            <div class='cap-desc'>Carregue documentos e faça perguntas precisas sobre o conteúdo via RAG</div>
        </div>
        <div class='cap-card'><div class='cap-icon'>🖼️</div>
            <div class='cap-title'>Visão Computacional</div>
            <div class='cap-desc'>Analise imagens, diagramas, fotos e prints com IA multimodal</div>
        </div>
        <div class='cap-card'><div class='cap-icon'>🌐</div>
            <div class='cap-title'>Busca em Tempo Real</div>
            <div class='cap-desc'>Respostas baseadas em informações atuais da internet</div>
        </div>
        <div class='cap-card'><div class='cap-icon'>🔊</div>
            <div class='cap-title'>Voz (TTS)</div>
            <div class='cap-desc'>Ouça as respostas em áudio com síntese de voz em português</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

# ── Input ────────────────────────────────────────────────────────────────────
uploaded_img = st.file_uploader(
    "📎 Imagem",
    type=["png","jpg","jpeg","webp"],
    key=f"img_{st.session_state.img_key}",
    label_visibility="collapsed",
)

prompt = st.chat_input("Pergunte qualquer coisa…")

# ═══════════════════════════════════════════════════════════════════════════════
#  PROCESSAMENTO
# ═══════════════════════════════════════════════════════════════════════════════
if prompt:
    # ── Imagem ───────────────────────────────────────────────────────────────
    img_path   = None
    img_b64    = None
    img_mime   = "image/jpeg"

    if uploaded_img:
        try:
            data = uploaded_img.read()
            ext  = uploaded_img.name.rsplit(".",1)[-1].lower()
            img_mime = f"image/{'jpeg' if ext=='jpg' else ext}"
            img_b64  = base64.b64encode(data).decode()
            fname    = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.{ext}"
            img_path = f"data/images/{fname}"
            with open(img_path,"wb") as f: f.write(data)
        except Exception as e:
            st.error(f"Erro ao processar imagem: {e}")

    # ── Salvar user msg ───────────────────────────────────────────────────────
    user_msg = {"role":"user","content":prompt}
    if img_path: user_msg["image"] = img_path
    st.session_state.messages.append(user_msg)

    db.ex(
        "INSERT INTO messages (conversation_id,role,content,image_path,model) VALUES (?,?,?,?,?)",
        (st.session_state.conv_id,"user",prompt,img_path,MODEL_ID)
    )

    # Título automático (primeira msg)
    if len(st.session_state.messages) == 1:
        db.ex(
            "UPDATE conversations SET title=?,updated_at=datetime('now','localtime') WHERE id=?",
            (prompt[:60].replace("\n"," "), st.session_state.conv_id)
        )

    # ── Exibir user msg ───────────────────────────────────────────────────────
    with st.chat_message("user"):
        if img_path and os.path.exists(img_path):
            st.image(img_path, width=380)
        st.markdown(prompt)

    # ── Contexto RAG ─────────────────────────────────────────────────────────
    rag_ctx     = ""
    rag_results = []
    web_ctx     = ""
    badges      = []

    if rag.count > 0 and not img_b64:
        with st.spinner("🔍 Buscando nos documentos…"):
            rag_ctx, rag_results = rag.search(prompt)
        if rag_ctx:
            badges.append(("📚 RAG", "badge-green"))
            log.info(f"RAG: {len(rag_results)} chunks encontrados")
        else:
            log.warning("RAG: nenhum resultado")

    # ── Contexto Web ─────────────────────────────────────────────────────────
    if st.session_state.web_on and not img_b64:
        with st.spinner("🌐 Buscando na web…"):
            web_ctx = web_search(prompt)
        if web_ctx: badges.append(("🌐 Web", "badge-teal"))

    if img_b64: badges.append(("🖼️ Visão", "badge-amber"))

    # ── Resposta ─────────────────────────────────────────────────────────────
    with st.chat_message("assistant"):

        # Badges de contexto
        if badges:
            bhtml = "".join(f"<span class='badge {cls}'>{lbl}</span>" for lbl,cls in badges)
            st.markdown(f"<div class='ctx-row'>{bhtml}</div>", unsafe_allow_html=True)

        answer     = ""
        used_model = MODEL_ID

        if img_b64:
            # Modo visão
            with st.spinner("🧠 Analisando imagem…"):
                answer, used_model = llm.vision(img_b64, prompt, img_mime,
                                                st.session_state.temp, st.session_state.maxtok)
            st.markdown(answer)

        else:
            # Modo texto
            system = build_system(rag_ctx, web_ctx)

            # Histórico: pegar mensagens sem duplicar a atual
            history = [
                {"role": m["role"], "content": m["content"]}
                for m in st.session_state.messages
            ]
            final = [{"role":"system","content":system}] + history

            if st.session_state.streaming:
                ph = st.empty()
                full = ""
                try:
                    for chunk in llm.stream(final, MODEL_ID,
                                            st.session_state.temp, st.session_state.maxtok):
                        full += chunk
                        ph.markdown(full + "▌")
                    ph.markdown(full)
                    answer = full
                except Exception as e:
                    answer = f"❌ Erro: {e}"
                    ph.markdown(answer)
                    used_model = "error"
            else:
                with st.spinner("💭 Pensando…"):
                    answer, used_model = llm.gen(final, MODEL_ID,
                                                 st.session_state.temp, st.session_state.maxtok)
                st.markdown(answer)

        # Footer
        st.markdown(
            f"<div class='model-footer'>"
            f"<span class='model-footer-tag'>🤖 {fmt_model(used_model)}</span>"
            f"</div>",
            unsafe_allow_html=True
        )

        # Debug RAG
        if st.session_state.show_debug and rag_results:
            with st.expander(f"🔍 Debug RAG — {len(rag_results)} chunks usados", expanded=False):
                for r in rag_results:
                    st.markdown(
                        f"**#{r['rank']} · {r['src']} · sim={r['sim']:.3f}**\n\n"
                        f"```\n{r['text'][:400]}\n```"
                    )

        # TTS
        if st.session_state.tts_on and answer and used_model != "error":
            with st.spinner("🔊 Gerando áudio…"):
                ap = tts(answer)
            if ap and os.path.exists(ap):
                with open(ap,"rb") as af:
                    st.audio(af.read(), format="audio/mp3")

        # Downloads
        c1,c2 = st.columns(2)
        c1.download_button("⬇️ TXT",
            export([{"role":"assistant","content":answer}],"txt"),
            "resposta.txt","text/plain")
        c2.download_button("⬇️ DOCX",
            export([{"role":"assistant","content":answer}],"docx"),
            "resposta.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # ── Salvar assistant msg ──────────────────────────────────────────────────
    st.session_state.messages.append({"role":"assistant","content":answer})
    db.ex(
        "INSERT INTO messages (conversation_id,role,content,model) VALUES (?,?,?,?)",
        (st.session_state.conv_id,"assistant",answer,used_model)
    )
    db.ex(
        "UPDATE conversations SET updated_at=datetime('now','localtime') WHERE id=?",
        (st.session_state.conv_id,)
    )

    st.session_state.img_key += 1

# ─── CLEANUP ──────────────────────────────────────────────────────────────────
@atexit.register
def _bye():
    if "db" in st.session_state:
        try: st.session_state.db.close()
        except: pass
